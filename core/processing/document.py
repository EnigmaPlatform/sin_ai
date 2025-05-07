import io
import logging
from typing import Union, Optional
from pathlib import Path
import pdfminer.high_level
import docx
from PIL import Image
import pytesseract

class DocumentProcessor:
    def __init__(self):
        self.logger = logging.getLogger("DocumentProcessor")
        self.temp_dir = Path("temp/docs")
        self.temp_dir.mkdir(parents=True, exist_ok=True)

    def process_file(self, file_path: Union[str, Path], file_type: Optional[str] = None) -> dict:
        """Обработка документа с автоматическим определением типа"""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
            
        if not file_type:
            file_type = self._detect_file_type(path)
            
        try:
            if file_type == "pdf":
                return self._process_pdf(path)
            elif file_type == "docx":
                return self._process_docx(path)
            elif file_type == "image":
                return self._process_image(path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            self.logger.error(f"Failed to process {file_path}: {str(e)}")
            raise

    def _detect_file_type(self, path: Path) -> str:
        """Определение типа файла по расширению"""
        ext = path.suffix.lower()
        if ext == ".pdf":
            return "pdf"
        elif ext == ".docx":
            return "docx"
        elif ext in (".jpg", ".jpeg", ".png", ".bmp"):
            return "image"
        else:
            raise ValueError(f"Unsupported file extension: {ext}")

    def _process_pdf(self, path: Path) -> dict:
        """Извлечение текста из PDF"""
        text = pdfminer.high_level.extract_text(path)
        return {
            "type": "pdf",
            "text": text,
            "pages": len(text.split("\f")),
            "metadata": self._get_pdf_metadata(path)
        }

    def _process_docx(self, path: Path) -> dict:
        """Извлечение текста из DOCX"""
        doc = docx.Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables = []
        
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
            
        return {
            "type": "docx",
            "paragraphs": paragraphs,
            "tables": tables,
            "metadata": {
                "author": doc.core_properties.author,
                "created": doc.core_properties.created.isoformat() if doc.core_properties.created else None
            }
        }

    def _process_image(self, path: Path) -> dict:
        """Извлечение текста из изображения"""
        img = Image.open(path)
        text = pytesseract.image_to_string(img)
        
        return {
            "type": "image",
            "text": text,
            "dimensions": f"{img.width}x{img.height}",
            "format": img.format
        }

    def _get_pdf_metadata(self, path: Path) -> dict:
        """Получение метаданных PDF"""
        from pdfminer.pdfparser import PDFParser
        from pdfminer.pdfdocument import PDFDocument
        
        with open(path, 'rb') as f:
            parser = PDFParser(f)
            doc = PDFDocument(parser)
            
            return {
                "title": doc.info[0].get('Title', b'').decode('utf-8', errors='ignore'),
                "author": doc.info[0].get('Author', b'').decode('utf-8', errors='ignore'),
                "pages": len(list(doc.get_pages()))
            }
